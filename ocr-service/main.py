"""
FastAPI OCR Service
Microservice for handling OCR processing with Surya, PaddleOCR, and Hybrid modes
Supports images (JPG, PNG, BMP, TIFF), PDF, and DOCX files
"""
from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
import uvicorn
import logging
import sys
import os
import time
from pathlib import Path
from typing import Optional, List
import shutil
import io
os.environ["PYTORCH_NO_CUDA_MEMORY_CACHING"] = "1"

# PDF and DOCX support
try:
    from pdf2image import convert_from_path, convert_from_bytes
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    
try:
    from docx import Document as DocxDocument
    from PIL import Image
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
# Configure logging for production
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(name)s: %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)

logger = logging.getLogger("ocr-service")

# Import OCR processing functions
sys.path.insert(0, str(Path(__file__).parent))
from modules.surya_utils import perform_ocr as perform_surya_ocr
from modules.paddle_utils import perform_paddle_ocr, perform_hybrid_ocr
from modules.preprocessing import preprocess_image, image_to_base64
import cv2
import numpy as np

app = FastAPI(
    title="OCR Service",
    description="Microservice for OCR processing with multiple engines",
    version="1.0.0"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Create upload directory
UPLOAD_DIR = Path("/app/uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# Supported file types
IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif'}
DOCUMENT_EXTENSIONS = {'.pdf', '.doc', '.docx'}


def convert_pdf_to_images(pdf_path: Path) -> List[np.ndarray]:
    """Convert PDF pages to list of OpenCV images"""
    if not PDF_SUPPORT:
        raise HTTPException(status_code=500, detail="PDF support not available. Please install pdf2image and poppler.")
    
    try:
        # Convert PDF to PIL images
        pil_images = convert_from_path(str(pdf_path), dpi=300)
        
        # Convert PIL images to OpenCV format (BGR)
        cv_images = []
        for pil_img in pil_images:
            # Convert PIL RGB to OpenCV BGR
            cv_img = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)
            cv_images.append(cv_img)
        
        return cv_images
    except Exception as e:
        logger.error(f"Failed to convert PDF: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to convert PDF: {str(e)}")


def extract_text_from_docx(docx_path: Path) -> str:
    """Extract text directly from DOCX file"""
    if not DOCX_SUPPORT:
        raise HTTPException(status_code=500, detail="DOCX support not available. Please install python-docx.")
    
    try:
        doc = DocxDocument(str(docx_path))
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(' | '.join(row_text))
        
        return '\n'.join(full_text)
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to extract text from DOCX: {str(e)}")


def get_file_type(filename: str, content_type: str) -> str:
    """Determine file type from filename and content type"""
    ext = Path(filename).suffix.lower() if filename else ''
    
    if ext in IMAGE_EXTENSIONS:
        return 'image'
    elif ext == '.pdf' or content_type == 'application/pdf':
        return 'pdf'
    elif ext in {'.doc', '.docx'} or content_type in {
        'application/msword',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    }:
        return 'docx'
    elif content_type and content_type.startswith('image/'):
        return 'image'
    
    return 'unknown'


def extract_text_from_results(ocr_results):
    """Extract plain text from OCR results, sorted by position"""
    if not ocr_results or 'text_lines' not in ocr_results:
        return ""

    text_lines = []
    for line in ocr_results['text_lines']:
        if 'text' in line and 'bbox' in line:
            bbox = line['bbox']
            y_pos = bbox[1] if len(bbox) >= 2 else 0
            x_pos = bbox[0] if len(bbox) >= 1 else 0
            text_lines.append({
                'text': line['text'],
                'y': y_pos,
                'x': x_pos
            })

    text_lines.sort(key=lambda item: (item['y'], item['x']))

    if not text_lines:
        return ""

    grouped_lines = []
    current_line_group = []
    current_y = text_lines[0]['y']
    y_threshold = 20

    for item in text_lines:
        if abs(item['y'] - current_y) < y_threshold:
            current_line_group.append(item['text'])
        else:
            if current_line_group:
                grouped_lines.append(' '.join(current_line_group))
            current_line_group = [item['text']]
            current_y = item['y']

    if current_line_group:
        grouped_lines.append(' '.join(current_line_group))

    return '\n'.join(grouped_lines)


def clean_results_for_json(results):
    """Clean OCR results for JSON serialization"""
    if not results:
        return None

    cleaned = {
        'layout': {'regions': []},
        'tables': [],
        'text_lines': []
    }

    if 'layout' in results and 'regions' in results['layout']:
        for region in results['layout']['regions']:
            cleaned['layout']['regions'].append({
                'bbox': region.get('bbox', []),
                'type': region.get('type', 'Unknown')
            })

    if 'tables' in results:
        for table in results['tables']:
            cleaned['tables'].append({
                'bbox': table.get('bbox', []),
                'confidence': float(table.get('confidence', 1.0))
            })

    if 'text_lines' in results:
        for line in results['text_lines']:
            cleaned['text_lines'].append({
                'text': line.get('text', ''),
                'bbox': line.get('bbox', []),
                'confidence': float(line.get('confidence', 1.0)),
                'region_type': line.get('region_type', 'Unknown')
            })

    return cleaned


@app.get("/")
async def root():
    """Health check endpoint"""
    return {
        "service": "OCR Service",
        "status": "healthy",
        "version": "1.0.0"
    }


@app.get("/health")
async def health():
    """Health check endpoint"""
    return {"status": "healthy"}


@app.post("/ocr/process")
async def process_ocr(
    file: UploadFile = File(...),
    preprocessing: bool = Form(default=True),
    engine: str = Form(default="hybrid")
):
    """
    Process OCR on uploaded file

    Args:
        file: Image file (JPEG, PNG, etc.), PDF, or DOCX
        preprocessing: Apply preprocessing (default: True)
        engine: OCR engine to use - 'surya', 'paddle', or 'hybrid' (default: 'hybrid')

    Returns:
        JSON response with OCR results
    """
    request_id = f"{int(time.time() * 1000)}"
    logger.info(f"[REQ {request_id}] ========== OCR REQUEST START ==========")
    start_time = time.time()

    temp_file_path = None

    try:
        # Determine file type
        file_type = get_file_type(file.filename, file.content_type)
        logger.info(f"[REQ {request_id}] File: {file.filename}, Type: {file_type}, Content-Type: {file.content_type}")
        
        if file_type == 'unknown':
            logger.error(f"[REQ {request_id}] Unsupported file type: {file.content_type}")
            raise HTTPException(status_code=400, detail="Unsupported file type. Please upload an image (JPG, PNG, BMP, TIFF), PDF, or DOCX file.")

        # Validate engine
        if engine not in ['surya', 'paddle', 'hybrid']:
            logger.error(f"[REQ {request_id}] Invalid engine: {engine}")
            raise HTTPException(status_code=400, detail="Engine must be 'surya', 'paddle', or 'hybrid'")

        logger.info(f"[REQ {request_id}] Engine: {engine}")
        logger.info(f"[REQ {request_id}] Preprocessing: {preprocessing}")

        # Save uploaded file temporarily
        save_start = time.time()
        temp_file_path = UPLOAD_DIR / f"{request_id}_{file.filename}"
        with temp_file_path.open("wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        file_size_kb = temp_file_path.stat().st_size / 1024
        logger.info(f"[REQ {request_id}] File saved: {file_size_kb:.2f} KB (took {time.time() - save_start:.2f}s)")

        result = {
            'text': '',
            'preprocessedImage': None,
            'preprocessingMetadata': None,
            'results': None
        }

        # Handle DOCX files - extract text directly
        if file_type == 'docx':
            logger.info(f"[REQ {request_id}] Processing DOCX file - extracting text directly")
            text_content = extract_text_from_docx(temp_file_path)
            
            result['text'] = text_content
            result['results'] = {
                'layout': {'regions': []},
                'tables': [],
                'text_lines': [{
                    'text': line,
                    'bbox': [0, i * 20, 100, (i + 1) * 20],
                    'confidence': 1.0,
                    'region_type': 'Text'
                } for i, line in enumerate(text_content.split('\n')) if line.strip()]
            }
            
            total_duration = time.time() - start_time
            logger.info(f"[REQ {request_id}] DOCX text extraction completed: {len(text_content)} characters")
            logger.info(f"[REQ {request_id}] ========== REQUEST COMPLETED in {total_duration:.2f}s ==========")
            
            return JSONResponse(content=result)

        # Handle PDF and Image files
        images_to_process = []
        
        if file_type == 'pdf':
            logger.info(f"[REQ {request_id}] Converting PDF to images...")
            convert_start = time.time()
            images_to_process = convert_pdf_to_images(temp_file_path)
            logger.info(f"[REQ {request_id}] PDF converted to {len(images_to_process)} page(s) (took {time.time() - convert_start:.2f}s)")
        else:
            # Load image
            load_start = time.time()
            image = cv2.imread(str(temp_file_path))
            if image is None:
                logger.error(f"[REQ {request_id}] Failed to load image")
                raise HTTPException(status_code=400, detail="Failed to load image")
            images_to_process = [image]
            logger.info(f"[REQ {request_id}] Image loaded: {image.shape} (took {time.time() - load_start:.2f}s)")

        # Process all images (pages)
        all_text_lines = []
        all_text_content = []
        preprocessed_image_b64 = None
        preprocessing_metadata = None
        
        for page_idx, image in enumerate(images_to_process):
            logger.info(f"[REQ {request_id}] Processing page {page_idx + 1}/{len(images_to_process)}")
            
            # Preprocessing
            image_for_ocr = image
            if preprocessing:
                logger.info(f"[REQ {request_id}] Starting preprocessing for page {page_idx + 1}...")
                preprocess_start = time.time()

                preprocessed, metadata = preprocess_image(image, save_steps_dir=None)
                image_for_ocr = preprocessed

                preprocess_duration = time.time() - preprocess_start
                logger.info(f"[REQ {request_id}] Preprocessing completed in {preprocess_duration:.2f}s")
                
                # Only store preprocessed image for first page (or single image)
                if page_idx == 0:
                    b64_start = time.time()
                    preprocessed_image_b64 = image_to_base64(preprocessed)
                    preprocessing_metadata = metadata
                    logger.info(f"[REQ {request_id}] Base64 encoding took {time.time() - b64_start:.2f}s")

            # OCR Processing
            logger.info(f"[REQ {request_id}] Starting OCR with engine: {engine} for page {page_idx + 1}")
            ocr_start = time.time()

            ocr_results = None
            if engine == 'surya':
                ocr_results = perform_surya_ocr(image_for_ocr, use_cuda=True)
            elif engine == 'paddle':
                ocr_results = perform_paddle_ocr(image_for_ocr, use_cuda=True)
            else:  # hybrid
                ocr_results = perform_hybrid_ocr(image_for_ocr, use_cuda=True)

            ocr_duration = time.time() - ocr_start
            logger.info(f"[REQ {request_id}] OCR completed for page {page_idx + 1} in {ocr_duration:.2f}s")

            if ocr_results and 'text_lines' in ocr_results:
                # Add page offset to bbox for multi-page documents
                page_offset = page_idx * 1000  # Arbitrary offset to separate pages
                for line in ocr_results['text_lines']:
                    if 'bbox' in line and len(line['bbox']) >= 2:
                        line['bbox'][1] += page_offset  # Offset y position
                        if len(line['bbox']) >= 4:
                            line['bbox'][3] += page_offset
                    line['page'] = page_idx + 1
                    all_text_lines.extend([line])
                
                # Extract text from this page
                page_text = extract_text_from_results(ocr_results)
                if page_text:
                    all_text_content.append(f"--- Page {page_idx + 1} ---\n{page_text}" if len(images_to_process) > 1 else page_text)

        # Combine results
        if not all_text_lines and not all_text_content:
            logger.warning(f"[REQ {request_id}] OCR processing returned no results")
            raise HTTPException(status_code=500, detail="OCR processing returned no results")

        result['text'] = '\n\n'.join(all_text_content)
        result['preprocessedImage'] = preprocessed_image_b64
        result['preprocessingMetadata'] = preprocessing_metadata
        result['results'] = {
            'layout': {'regions': []},
            'tables': [],
            'text_lines': [{
                'text': line.get('text', ''),
                'bbox': line.get('bbox', []),
                'confidence': float(line.get('confidence', 1.0)),
                'region_type': line.get('region_type', 'Unknown'),
                'page': line.get('page', 1)
            } for line in all_text_lines]
        }

        total_duration = time.time() - start_time
        logger.info(f"[REQ {request_id}] Extracted {len(result['text'])} characters from {len(images_to_process)} page(s)")
        logger.info(f"[REQ {request_id}] ========== REQUEST COMPLETED in {total_duration:.2f}s ==========")

        return JSONResponse(content=result)

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[REQ {request_id}] OCR processing failed: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"OCR processing failed: {str(e)}")

    finally:
        # Cleanup temporary file
        if temp_file_path and temp_file_path.exists():
            try:
                temp_file_path.unlink()
                logger.info(f"[REQ {request_id}] Temporary file cleaned up")
            except Exception as e:
                logger.warning(f"[REQ {request_id}] Failed to cleanup temp file: {e}")


if __name__ == "__main__":
    logger.info("Starting OCR Service on port 7000...")
    uvicorn.run(app, host="0.0.0.0", port=7000)
