"""
RunPod Serverless handler for VLM-based OCR (Qwen2.5-VL-7B).

Supports engine: vlm only
Handles images, PDFs, and DOCX files.

Request format:
{
    "input": {
        "image": "<base64-encoded file>",
        "filename": "document.pdf",
        "preprocessing": false
    }
}
"""

import os
import sys
import time
import base64
import tempfile
import logging

import runpod
import cv2
import numpy as np

sys.path.insert(0, '/app')

from vlm_engine import VLMEngine

# PDF support
try:
    from pdf2image import convert_from_path
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

# DOCX support
try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')
logger = logging.getLogger("runpod-vlm")

# Lazy-loaded VLM engine singleton
_vlm_engine = None

IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif'}


def get_vlm_engine():
    """Lazy-load VLM engine on first use."""
    global _vlm_engine
    if _vlm_engine is None:
        logger.info("Initializing VLM engine (first request)...")
        _vlm_engine = VLMEngine()
        logger.info("VLM engine ready!")
    return _vlm_engine


def get_file_type(filename: str) -> str:
    if not filename:
        return 'image'
    ext = os.path.splitext(filename)[1].lower()
    if ext in IMAGE_EXTENSIONS:
        return 'image'
    elif ext == '.pdf':
        return 'pdf'
    elif ext in {'.doc', '.docx'}:
        return 'docx'
    return 'image'


def convert_pdf_to_images(pdf_path: str):
    if not PDF_SUPPORT:
        raise RuntimeError("PDF support not available. pdf2image/poppler not installed.")
    pil_images = convert_from_path(pdf_path, dpi=300)
    return [cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR) for img in pil_images]


def extract_text_from_docx(docx_path: str) -> str:
    if not DOCX_SUPPORT:
        raise RuntimeError("DOCX support not available. python-docx not installed.")
    doc = DocxDocument(docx_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(' | '.join(row_text))
    return '\n'.join(full_text)


def process_vlm(input_data):
    """Main VLM processing function."""
    start_time = time.time()

    # Parse input
    image_b64 = input_data.get("image")
    image_url = input_data.get("image_url")
    filename = input_data.get("filename", "image.jpg")
    custom_prompt = input_data.get("prompt")  # Optional custom prompt

    if not image_b64 and not image_url:
        return {"error": "Provide 'image' (base64) or 'image_url' in input"}

    logger.info(f"Processing: filename={filename}")

    # Decode file
    try:
        if image_url and not image_b64:
            import requests
            resp = requests.get(image_url, timeout=30)
            resp.raise_for_status()
            file_bytes = resp.content
            image_b64 = base64.b64encode(file_bytes).decode('utf-8')
        else:
            file_bytes = base64.b64decode(image_b64)
    except Exception as e:
        return {"error": f"Failed to decode input: {str(e)}"}

    # Save to temp file (needed for PDF/DOCX)
    tmp_dir = tempfile.mkdtemp()
    tmp_path = os.path.join(tmp_dir, filename)
    with open(tmp_path, 'wb') as f:
        f.write(file_bytes)

    try:
        file_type = get_file_type(filename)
        logger.info(f"File type: {file_type}, size: {len(file_bytes) / 1024:.1f} KB")

        vlm = get_vlm_engine()

        result = {
            'text': '',
            'preprocessedImage': None,
            'preprocessingMetadata': None,
            'results': None
        }

        # Handle DOCX - direct text extraction (no need for VLM)
        if file_type == 'docx':
            text_content = extract_text_from_docx(tmp_path)
            result['text'] = text_content
            result['results'] = {
                'layout': {'regions': []},
                'tables': [],
                'text_lines': [{
                    'text': line,
                    'bbox': [0, i * 20, 100, (i + 1) * 20],
                    'confidence': 1.0,
                    'region_type': 'Text'
                } for i, line in enumerate(text_content.split('\n')) if line.strip()]
            }
            result['_timing_ms'] = int((time.time() - start_time) * 1000)
            return result

        # Handle PDF - convert pages to images for VLM
        if file_type == 'pdf':
            images = convert_pdf_to_images(tmp_path)
            all_text = []
            all_text_lines = []

            for page_idx, cv_img in enumerate(images):
                _, buf = cv2.imencode('.jpg', cv_img)
                page_b64 = base64.b64encode(buf).decode('utf-8')

                vlm_result = vlm.process_image(page_b64, prompt=custom_prompt)
                page_text = vlm_result.get('text', '')
                if page_text:
                    prefix = f"--- Page {page_idx + 1} ---\n" if len(images) > 1 else ""
                    all_text.append(f"{prefix}{page_text}")

                for line in vlm_result.get('text_lines', []):
                    line['page'] = page_idx + 1
                    all_text_lines.append(line)

            result['text'] = '\n\n'.join(all_text)
            result['results'] = {
                'layout': {'regions': []},
                'tables': [],
                'text_lines': all_text_lines
            }
        else:
            # Single image
            vlm_result = vlm.process_image(image_b64, prompt=custom_prompt)
            result['text'] = vlm_result.get('text', '')
            result['results'] = {
                'layout': {'regions': []},
                'tables': [],
                'text_lines': vlm_result.get('text_lines', [])
            }
            if '_vlm_metadata' in vlm_result:
                result['_vlm_metadata'] = vlm_result['_vlm_metadata']

        result['_timing_ms'] = int((time.time() - start_time) * 1000)
        return result

    except Exception as e:
        logger.error(f"Processing failed: {e}", exc_info=True)
        return {"error": str(e)}
    finally:
        import shutil
        try:
            shutil.rmtree(tmp_dir)
        except:
            pass


def handler(event):
    """RunPod serverless handler entry point."""
    input_data = event.get("input", {})
    return process_vlm(input_data)


# Startup
print("=" * 60)
print("VLM Service - RunPod Serverless Handler")
print(f"Model: {os.environ.get('MODEL_NAME', 'Qwen/Qwen2.5-VL-7B-Instruct')}")
print(f"PDF support: {PDF_SUPPORT}")
print(f"DOCX support: {DOCX_SUPPORT}")
print("Engine: vlm (Qwen2.5-VL-7B via vLLM)")
print("=" * 60)

runpod.serverless.start({"handler": handler})
