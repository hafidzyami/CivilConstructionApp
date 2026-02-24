"""
RunPod Serverless handler for OCR processing (no VLM).

Supports engines: surya, paddle, hybrid
Handles images, PDFs, and DOCX files.

Request format:
{
    "input": {
        "image": "<base64-encoded file>",
        "filename": "document.pdf",
        "engine": "hybrid",
        "preprocessing": true
    }
}
"""

import os
import sys
import time
import base64
import tempfile
import gc
import logging

import runpod
import cv2
import numpy as np

# Add parent directory so we can import modules/
sys.path.insert(0, '/app')

from modules.surya_utils import perform_ocr as perform_surya_ocr
from modules.paddle_utils import perform_paddle_ocr, perform_hybrid_ocr
from modules.preprocessing import preprocess_image, image_to_base64
from modules.device_utils import get_device

# PDF and DOCX support
try:
    from pdf2image import convert_from_path
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')
logger = logging.getLogger("runpod-ocr")

# Supported file types
IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif'}
DOCUMENT_EXTENSIONS = {'.pdf', '.doc', '.docx'}


def get_file_type(filename: str) -> str:
    """Determine file type from filename."""
    if not filename:
        return 'image'
    ext = os.path.splitext(filename)[1].lower()
    if ext in IMAGE_EXTENSIONS:
        return 'image'
    elif ext == '.pdf':
        return 'pdf'
    elif ext in {'.doc', '.docx'}:
        return 'docx'
    return 'image'


def convert_pdf_to_images(pdf_path: str):
    """Convert PDF pages to list of OpenCV images."""
    if not PDF_SUPPORT:
        raise RuntimeError("PDF support not available. pdf2image/poppler not installed.")
    pil_images = convert_from_path(pdf_path, dpi=300)
    return [cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR) for img in pil_images]


def extract_text_from_docx(docx_path: str) -> str:
    """Extract text directly from DOCX file."""
    if not DOCX_SUPPORT:
        raise RuntimeError("DOCX support not available. python-docx not installed.")
    doc = DocxDocument(docx_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(' | '.join(row_text))
    return '\n'.join(full_text)


def extract_text_from_results(ocr_results):
    """Extract plain text from OCR results, sorted by position."""
    if not ocr_results or 'text_lines' not in ocr_results:
        return ""

    text_lines = []
    for line in ocr_results['text_lines']:
        if 'text' in line and 'bbox' in line:
            bbox = line['bbox']
            text_lines.append({
                'text': line['text'],
                'y': bbox[1] if len(bbox) >= 2 else 0,
                'x': bbox[0] if len(bbox) >= 1 else 0
            })

    text_lines.sort(key=lambda item: (item['y'], item['x']))
    if not text_lines:
        return ""

    grouped_lines = []
    current_group = []
    current_y = text_lines[0]['y']
    y_threshold = 20

    for item in text_lines:
        if abs(item['y'] - current_y) < y_threshold:
            current_group.append(item['text'])
        else:
            if current_group:
                grouped_lines.append(' '.join(current_group))
            current_group = [item['text']]
            current_y = item['y']

    if current_group:
        grouped_lines.append(' '.join(current_group))

    return '\n'.join(grouped_lines)


def process_ocr(input_data):
    """Main OCR processing function."""
    start_time = time.time()

    # Parse input
    image_b64 = input_data.get("image")
    image_url = input_data.get("image_url")
    filename = input_data.get("filename", "image.jpg")
    engine = input_data.get("engine", "hybrid")
    do_preprocessing = input_data.get("preprocessing", True)

    if not image_b64 and not image_url:
        return {"error": "Provide 'image' (base64) or 'image_url' in input"}

    if engine not in ['surya', 'paddle', 'hybrid']:
        return {"error": f"Invalid engine: {engine}. Must be surya, paddle, or hybrid"}

    logger.info(f"Processing: engine={engine}, filename={filename}, preprocessing={do_preprocessing}")

    # Decode file
    try:
        if image_url and not image_b64:
            import requests
            resp = requests.get(image_url, timeout=30)
            resp.raise_for_status()
            file_bytes = resp.content
        else:
            file_bytes = base64.b64decode(image_b64)
    except Exception as e:
        return {"error": f"Failed to decode input: {str(e)}"}

    # Save to temp file
    tmp_dir = tempfile.mkdtemp()
    tmp_path = os.path.join(tmp_dir, filename)
    with open(tmp_path, 'wb') as f:
        f.write(file_bytes)

    try:
        file_type = get_file_type(filename)
        logger.info(f"File type: {file_type}, size: {len(file_bytes) / 1024:.1f} KB")

        result = {
            'text': '',
            'preprocessedImage': None,
            'preprocessingMetadata': None,
            'results': None
        }

        # Handle DOCX - direct text extraction
        if file_type == 'docx':
            text_content = extract_text_from_docx(tmp_path)
            result['text'] = text_content
            result['results'] = {
                'layout': {'regions': []},
                'tables': [],
                'text_lines': [{
                    'text': line,
                    'bbox': [0, i * 20, 100, (i + 1) * 20],
                    'confidence': 1.0,
                    'region_type': 'Text'
                } for i, line in enumerate(text_content.split('\n')) if line.strip()]
            }
            result['_timing_ms'] = int((time.time() - start_time) * 1000)
            return result

        # Handle images and PDFs with OCR engines
        images_to_process = []
        if file_type == 'pdf':
            logger.info("Converting PDF to images...")
            images_to_process = convert_pdf_to_images(tmp_path)
            logger.info(f"PDF converted to {len(images_to_process)} page(s)")
        else:
            image = cv2.imread(tmp_path)
            if image is None:
                return {"error": "Failed to load image"}
            images_to_process = [image]

        all_text_lines = []
        all_text_content = []
        preprocessed_image_b64 = None
        all_preprocessed_images = []
        preprocessing_metadata = None

        for page_idx, image in enumerate(images_to_process):
            logger.info(f"Processing page {page_idx + 1}/{len(images_to_process)}")

            # Preprocessing
            image_for_ocr = image
            if do_preprocessing:
                preprocess_start = time.time()
                preprocessed, metadata = preprocess_image(image, save_steps_dir=None)
                image_for_ocr = preprocessed
                logger.info(f"Preprocessing took {time.time() - preprocess_start:.2f}s")

                page_preprocessed_b64 = image_to_base64(preprocessed)
                all_preprocessed_images.append(page_preprocessed_b64)
                if page_idx == 0:
                    preprocessed_image_b64 = page_preprocessed_b64
                    preprocessing_metadata = metadata

            # OCR
            ocr_start = time.time()
            ocr_results = None
            if engine == 'surya':
                ocr_results = perform_surya_ocr(image_for_ocr, use_cuda=True)
            elif engine == 'paddle':
                ocr_results = perform_paddle_ocr(image_for_ocr, use_cuda=True)
            else:  # hybrid
                ocr_results = perform_hybrid_ocr(image_for_ocr, use_cuda=True)

            logger.info(f"OCR ({engine}) took {time.time() - ocr_start:.2f}s")

            if ocr_results and 'text_lines' in ocr_results:
                page_offset = page_idx * 1000
                for line in ocr_results['text_lines']:
                    if 'bbox' in line and len(line['bbox']) >= 2:
                        line['bbox'][1] += page_offset
                        if len(line['bbox']) >= 4:
                            line['bbox'][3] += page_offset
                    line['page'] = page_idx + 1
                    all_text_lines.append(line)

                page_text = extract_text_from_results(ocr_results)
                if page_text:
                    prefix = f"--- Page {page_idx + 1} ---\n" if len(images_to_process) > 1 else ""
                    all_text_content.append(f"{prefix}{page_text}")

            # Free memory between pages
            gc.collect()

        if not all_text_lines and not all_text_content:
            return {"error": "OCR processing returned no results"}

        result['text'] = '\n\n'.join(all_text_content)
        result['preprocessedImage'] = preprocessed_image_b64
        result['preprocessedImages'] = all_preprocessed_images if len(all_preprocessed_images) > 1 else None
        result['preprocessingMetadata'] = preprocessing_metadata
        result['results'] = {
            'layout': {'regions': []},
            'tables': [],
            'text_lines': [{
                'text': line.get('text', ''),
                'bbox': line.get('bbox', []),
                'confidence': float(line.get('confidence', 1.0)),
                'region_type': line.get('region_type', 'Unknown'),
                'page': line.get('page', 1)
            } for line in all_text_lines]
        }

        result['_timing_ms'] = int((time.time() - start_time) * 1000)
        return result

    except Exception as e:
        logger.error(f"Processing failed: {e}", exc_info=True)
        return {"error": str(e)}
    finally:
        # Cleanup temp files
        import shutil
        try:
            shutil.rmtree(tmp_dir)
        except:
            pass


def handler(event):
    """RunPod serverless handler entry point."""
    input_data = event.get("input", {})
    return process_ocr(input_data)


# Startup
print("=" * 60)
print("OCR Service - RunPod Serverless Handler")
print(f"GPU: {get_device(prefer_cuda=True).upper()}")
print(f"PDF support: {PDF_SUPPORT}")
print(f"DOCX support: {DOCX_SUPPORT}")
print("Engines: surya, paddle, hybrid")
print("=" * 60)

runpod.serverless.start({"handler": handler})
