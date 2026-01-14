#!/usr/bin/env python3
"""
PDF/DOCX Regulation Extractor
Extracts building regulation articles and outputs to JSON format
"""

import re
import json
import sys
from pathlib import Path
from typing import List, Dict, Any, Optional
import pdfplumber
from docx import Document


class RegulationExtractor:
    """Extracts structured data from regulation documents"""

    def __init__(self):
        self.current_article = None
        self.articles = []
        self.sub_articles = []

    def extract_from_pdf(self, pdf_path: str) -> Dict[str, Any]:
        """Extract text and articles from PDF file"""
        print(f"[PDF] Extracting from: {pdf_path}")

        full_text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text:
                    full_text += text + "\n"
                print(f"   Page {page_num}/{len(pdf.pages)} extracted")

        return self._parse_articles(full_text)

    def extract_from_docx(self, docx_path: str) -> Dict[str, Any]:
        """Extract text and articles from DOCX file"""
        print(f"[DOCX] Extracting from: {docx_path}")

        doc = Document(docx_path)
        full_text = ""

        for para_num, para in enumerate(doc.paragraphs, 1):
            full_text += para.text + "\n"
            if para_num % 50 == 0:
                print(f"   {para_num} paragraphs processed...")

        print(f"   Total {len(doc.paragraphs)} paragraphs extracted")
        return self._parse_articles(full_text)

    def _parse_articles(self, text: str) -> Dict[str, Any]:
        """Parse articles from extracted text"""
        print("[INFO] Parsing articles...")

        # Clean and normalize text
        text = self._preprocess_text(text)

        # Split into lines
        lines = text.split('\n')

        articles = []
        sub_articles = []
        current_article_text = []
        current_article_id = None
        current_article_title = None

        # Patterns for matching articles
        # Matches: "Article 5 (Title)", "Article 27:", "제5조", "제27조"
        article_pattern = re.compile(r'^(?:Article\s+(\d+)\s*(?:\(([^)]+)\)|[:：]\s*(.*?))?|제\s*(\d+)\s*조\s*(.*?))', re.IGNORECASE)

        # Matches: "Article 5-1 (Title)", "Article 27-2:", "제5조의2"
        sub_article_pattern = re.compile(r'^(?:Article\s+(\d+)-(\d+)\s*(?:\(([^)]+)\)|[:：]\s*(.*?))?|제\s*(\d+)\s*조의\s*(\d+)\s*(.*?))', re.IGNORECASE)

        # Matches: "Article 5-1-1 (Title)"
        nested_sub_pattern = re.compile(r'^Article\s+(\d+(?:-\d+)+)\s*(?:\(([^)]+)\)|[:：]\s*(.*?))?', re.IGNORECASE)

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check for nested sub-article (e.g., 5-1-1)
            nested_match = nested_sub_pattern.match(line)
            if nested_match:
                # Save previous article
                if current_article_id and current_article_text:
                    self._save_article(articles, sub_articles, current_article_id,
                                     current_article_title, current_article_text)

                full_id = nested_match.group(1)
                title = nested_match.group(2) or nested_match.group(3) or ""
                title = title.strip()
                parts = full_id.split('-')
                level = len(parts) - 1
                parent_parts = parts[:-1]
                parent_id = '-'.join(parent_parts) if parent_parts else parts[0]

                current_article_id = f"article_{full_id}"
                current_article_title = title if title else f"Sub-Article {full_id}"
                current_article_text = []
                continue

            # Check for sub-article (e.g., 5-1, 27-2)
            sub_match = sub_article_pattern.match(line)
            if sub_match:
                # Save previous article
                if current_article_id and current_article_text:
                    self._save_article(articles, sub_articles, current_article_id,
                                     current_article_title, current_article_text)

                if sub_match.group(1):  # English format "Article X-Y"
                    base_num = sub_match.group(1)
                    sub_num = sub_match.group(2)
                    # Group 3 is title in parentheses, Group 4 is title after colon
                    title = (sub_match.group(3) or sub_match.group(4) or "").strip()
                else:  # Korean format "제X조의Y"
                    base_num = sub_match.group(5)
                    sub_num = sub_match.group(6)
                    title = (sub_match.group(7) or "").strip()

                current_article_id = f"article_{base_num}-{sub_num}"
                current_article_title = title if title else f"Sub-Article {base_num}-{sub_num}"
                current_article_text = []
                continue

            # Check for base article
            article_match = article_pattern.match(line)
            if article_match:
                # Save previous article
                if current_article_id and current_article_text:
                    self._save_article(articles, sub_articles, current_article_id,
                                     current_article_title, current_article_text)

                if article_match.group(1):  # English format "Article X"
                    article_num = article_match.group(1)
                    # Group 2 is title in parentheses, Group 3 is title after colon
                    title = (article_match.group(2) or article_match.group(3) or "").strip()
                else:  # Korean format "제X조"
                    article_num = article_match.group(4)
                    title = (article_match.group(5) or "").strip()

                current_article_id = f"article_{article_num}"
                current_article_title = title if title else f"Article {article_num}"
                current_article_text = []
                continue

            # Add line to current article
            if current_article_id:
                current_article_text.append(line)

        # Save last article
        if current_article_id and current_article_text:
            self._save_article(articles, sub_articles, current_article_id,
                             current_article_title, current_article_text)

        print(f"[SUCCESS] Extracted {len(articles)} base articles and {len(sub_articles)} sub-articles")

        return {
            'articles': articles,
            'sub_articles': sub_articles
        }

    def _save_article(self, articles: List, sub_articles: List,
                     article_id: str, title: str, text_lines: List[str]):
        """Save article to appropriate list"""
        text = ' '.join(text_lines).strip()

        # Determine if it's a base article or sub-article
        is_sub = '-' in article_id.replace('article_', '')

        article_data = {
            'id': article_id,
            'title': title,
            'text': text
        }

        if is_sub:
            # Parse level and parent
            number_part = article_id.replace('article_', '')
            parts = number_part.split('-')
            level = len(parts) - 1
            parent_parts = parts[:-1]
            parent_id = f"article_{'-'.join(parent_parts)}" if len(parent_parts) > 1 else f"article_{parent_parts[0]}"

            article_data['level'] = level
            article_data['parentId'] = parent_id

            sub_articles.append(article_data)
        else:
            articles.append(article_data)


    def _preprocess_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Normalize line breaks first
        text = text.replace('\r\n', '\n').replace('\r', '\n')

        # Remove excessive whitespace within lines (but keep newlines)
        lines = text.split('\n')
        lines = [re.sub(r'[ \t]+', ' ', line) for line in lines]
        text = '\n'.join(lines)

        # Remove page numbers
        text = re.sub(r'^\s*[-–—]*\s*\d+\s*[-–—]*\s*$', '', text, flags=re.MULTILINE)
        text = re.sub(r'^\s*Page\s+\d+\s*$', '', text, flags=re.MULTILINE | re.IGNORECASE)

        # Clean OCR errors
        text = text.replace(''', "'").replace(''', "'")
        text = text.replace('"', '"').replace('"', '"')

        # Normalize article headers
        text = re.sub(r'Article\s+(\d+(?:-\d+)*)\s*[:：]?\s*', r'Article \1: ', text, flags=re.IGNORECASE)
        text = re.sub(r'제\s*(\d+)\s*조\s*', r'Article \1: ', text)
        text = re.sub(r'제\s*(\d+)\s*조의\s*(\d+)', r'Article \1-\2: ', text)

        return text


def main():
    """Main extraction function"""
    print("=" * 60)
    print("Building Regulation Extractor")
    print("=" * 60)
    print()

    # Setup paths
    script_dir = Path(__file__).parent
    regulations_dir = script_dir.parent / "backend" / "data" / "regulations"
    output_dir = script_dir / "output"
    output_dir.mkdir(exist_ok=True)

    extractor = RegulationExtractor()

    # Extract National regulation (PDF - Enforcement Decree)
    national_path = regulations_dir / "enforcement-decree.pdf"
    if national_path.exists():
        print("[NATIONAL REGULATION]")
        print("-" * 60)
        national_data = extractor.extract_from_pdf(str(national_path))

        national_regulation = {
            "name": "Building Act (Enforcement Decree)",
            "level": "National",
            "authority": "Ministry of Land, Infrastructure and Transport",
            "articles": national_data['articles'],
            "subArticles": national_data['sub_articles']
        }

        output_path = output_dir / "national-regulation.json"
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(national_regulation, f, indent=2, ensure_ascii=False)

        print(f"[SAVED] {output_path}")
        print()
    else:
        print(f"[WARNING] National regulation not found: {national_path}")
        print()

    # Extract Regional regulation (DOCX - Building Act Sample)
    regional_path = regulations_dir / "building-act-sample.docx"
    if regional_path.exists():
        print("[REGIONAL REGULATION]")
        print("-" * 60)
        regional_data = extractor.extract_from_docx(str(regional_path))

        regional_regulation = {
            "name": "Seoul Building Ordinance",
            "level": "Regional",
            "authority": "Seoul Metropolitan Government",
            "articles": regional_data['articles'],
            "subArticles": regional_data['sub_articles']
        }

        output_path = output_dir / "regional-regulation.json"
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(regional_regulation, f, indent=2, ensure_ascii=False)

        print(f"[SAVED] {output_path}")
        print()
    else:
        print(f"[WARNING] Regional regulation not found: {regional_path}")
        print()

    print("=" * 60)
    print("[SUCCESS] Extraction complete!")
    print()
    print("Output files:")
    print(f"   - {output_dir / 'national-regulation.json'}")
    print(f"   - {output_dir / 'regional-regulation.json'}")
    print()
    print("Next steps:")
    print("   1. Review the extracted JSON files")
    print("   2. Manually add/correct any articles that were missed")
    print("   3. Add zone information if needed")
    print("   4. Copy the JSON files to backend/data/regulations/")
    print("=" * 60)


if __name__ == "__main__":
    main()
